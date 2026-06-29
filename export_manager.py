# -*- coding: utf-8 -*-
"""
내보내기 관리 모듈
- 엑셀(CSV) 내보내기
- 워드(HTML/DOCX) 내보내기
"""

from datetime import datetime
from qgis.core import Qgis, QgsMessageLog

from .constants import extract_jimok_from_jibun, extract_jimok_from_pnu


class ExportManager:
    """내보내기 관리 클래스 - 대시보드 및 모든 탭 정보 포함"""

    @staticmethod
    def export_table_xlsx(headers, rows, filepath):
        """단일 표를 엑셀(.xlsx)로 내보내기. openpyxl이 없으면 CSV(utf-8-sig)로 폴백.

        headers: [str, ...], rows: [[v, ...], ...]
        반환: 실제 저장된 파일 경로 (폴백 시 확장자 .csv로 변경됨)
        """
        try:
            import openpyxl

            wb = openpyxl.Workbook()
            ws = wb.active
            ws.append(list(headers))
            for row in rows:
                ws.append(['' if v is None else v for v in row])
            # 헤더 굵게 + 열 너비 자동(간이)
            for col_idx, header in enumerate(headers, start=1):
                cell = ws.cell(row=1, column=col_idx)
                cell.font = openpyxl.styles.Font(bold=True)
                width = max(
                    [len(str(header))]
                    + [len(str(r[col_idx - 1])) for r in rows
                       if len(r) >= col_idx]) if rows else len(str(header))
                ws.column_dimensions[
                    openpyxl.utils.get_column_letter(col_idx)].width = min(width + 4, 50)
            if not filepath.lower().endswith('.xlsx'):
                filepath += '.xlsx'
            wb.save(filepath)
            return filepath
        except ImportError:
            import csv

            csv_path = filepath
            if csv_path.lower().endswith('.xlsx'):
                csv_path = csv_path[:-5] + '.csv'
            elif not csv_path.lower().endswith('.csv'):
                csv_path += '.csv'
            with open(csv_path, 'w', newline='', encoding='utf-8-sig') as f:
                writer = csv.writer(f)
                writer.writerow(headers)
                writer.writerows(rows)
            QgsMessageLog.logMessage(
                "openpyxl 미설치 - CSV로 대체 저장", "VWorld", Qgis.Warning)
            return csv_path

    # ------------------------------------------------------------------
    # 분석 탭 요약 섹션 공통 렌더러 (v1.5.1)
    # 섹션 형식: {'title': str, 'kv': [(라벨, 값)], 'text': str,
    #            'tables': [{'title', 'headers', 'rows'}], 'images': [경로]}
    # ------------------------------------------------------------------
    @staticmethod
    def _write_sections_csv(writer, sections):
        for section in (sections or []):
            if not section:
                continue
            writer.writerow([f"=== {section.get('title', '분석')} ==="])
            for label, value in section.get('kv', []):
                writer.writerow([label, value])
            text = section.get('text')
            if text:
                for line in str(text).split('\n'):
                    writer.writerow([line])
            for table in section.get('tables', []):
                if table.get('title'):
                    writer.writerow([f"[{table['title']}]"])
                if table.get('headers'):
                    writer.writerow(table['headers'])
                for row in table.get('rows', []):
                    writer.writerow(['' if v is None else v for v in row])
                writer.writerow([])
            for path in section.get('images', []):
                writer.writerow(['이미지', path])
            writer.writerow([])

    @staticmethod
    def _write_sections_docx(doc, sections):
        for section in (sections or []):
            if not section:
                continue
            doc.add_heading(section.get('title', '분석'), level=1)
            for label, value in section.get('kv', []):
                doc.add_paragraph(f"{label}: {value}")
            text = section.get('text')
            if text:
                doc.add_paragraph(str(text)[:8000])
            for table_def in section.get('tables', []):
                headers = table_def.get('headers', [])
                rows = table_def.get('rows', [])
                if not headers and not rows:
                    continue
                if table_def.get('title'):
                    doc.add_heading(table_def['title'], level=2)
                max_rows = min(len(rows), 200)
                table = doc.add_table(rows=max_rows + 1, cols=max(len(headers), 1))
                table.style = 'Table Grid'
                for i, header in enumerate(headers):
                    table.rows[0].cells[i].text = str(header)
                for r, row in enumerate(rows[:max_rows], 1):
                    for c, value in enumerate(row[:len(headers) or len(row)]):
                        table.rows[r].cells[c].text = \
                            '' if value is None else str(value)
                doc.add_paragraph()
            for path in section.get('images', []):
                try:
                    from docx.shared import Inches
                    doc.add_picture(path, width=Inches(5.5))
                except Exception:
                    doc.add_paragraph(f"이미지: {path}")
            doc.add_paragraph()

    @staticmethod
    def _render_sections_html(sections):
        import html as html_mod

        def esc(v):
            return html_mod.escape('' if v is None else str(v))

        out = ""
        for section in (sections or []):
            if not section:
                continue
            out += f"\n    <h2>{esc(section.get('title', '분석'))}</h2>\n"
            kv = section.get('kv', [])
            if kv:
                out += "    <table>\n"
                for label, value in kv:
                    out += (f"        <tr><th style='width:30%'>{esc(label)}"
                            f"</th><td>{esc(value)}</td></tr>\n")
                out += "    </table>\n"
            text = section.get('text')
            if text:
                out += (f"    <div class=\"debug-log\">{esc(str(text)[:12000])}"
                        "</div>\n")
            for table_def in section.get('tables', []):
                if table_def.get('title'):
                    out += f"    <h3>{esc(table_def['title'])}</h3>\n"
                out += "    <table>\n        <tr>"
                for header in table_def.get('headers', []):
                    out += f"<th>{esc(header)}</th>"
                out += "</tr>\n"
                for row in table_def.get('rows', [])[:300]:
                    out += ("        <tr>"
                            + "".join(f"<td>{esc(v)}</td>" for v in row)
                            + "</tr>\n")
                out += "    </table>\n"
            for path in section.get('images', []):
                src = str(path).replace('\\', '/')
                out += (f"    <p><img src=\"file:///{esc(src)}\" "
                        f"style=\"max-width:480px\"> <br>{esc(path)}</p>\n")
        return out

    @staticmethod
    def export_to_excel(data, filepath, dashboard_stats=None, debug_log=None,
                        analysis_summaries=None):
        """엑셀 형식으로 내보내기 (대시보드, 모든 탭 데이터, 디버그 로그 포함)"""
        try:
            import csv

            with open(filepath, 'w', newline='', encoding='utf-8-sig') as f:
                writer = csv.writer(f)

                # 대시보드 요약 정보
                writer.writerow(['=== 대시보드 요약 ==='])
                writer.writerow([])

                if dashboard_stats:
                    summary = dashboard_stats.get('summary', {})
                    writer.writerow(['총 필지 수', summary.get('total_count', 0)])
                    writer.writerow(['총 토지면적(m2)', f"{summary.get('total_area', 0):,.2f}"])
                    writer.writerow(['총 편입면적(m2)', f"{summary.get('total_inclusion_area', 0):,.2f}"])
                    writer.writerow(['평균 공시지가(원/m2)', f"{summary.get('avg_price', 0):,.0f}"])
                    writer.writerow(['최고 공시지가(원/m2)', f"{summary.get('max_price', 0):,.0f}"])
                    writer.writerow(['최저 공시지가(원/m2)', f"{summary.get('min_price', 0):,.0f}"])
                    writer.writerow(['총 공시지가 합계(원)', f"{summary.get('total_price_sum', 0):,.0f}"])
                    writer.writerow(['토지 수', summary.get('land_count', 0)])
                    writer.writerow(['임야 수', summary.get('forest_count', 0)])
                    writer.writerow(['소유자 수', summary.get('owner_count', 0)])
                    writer.writerow([])

                    # 시도별 통계
                    writer.writerow(['=== 시도별 편입면적 분석 ==='])
                    writer.writerow(['시도', '필지 수', '토지면적(m2)', '편입면적(m2)', '공시지가합계(원)'])
                    for sido, stats in sorted(dashboard_stats.get('sido_stats', {}).items()):
                        writer.writerow([
                            sido, stats['count'],
                            f"{stats['area']:,.2f}", f"{stats['inclusion_area']:,.2f}",
                            f"{stats['price_sum']:,.0f}"
                        ])
                    writer.writerow([])

                    # 시군구별 통계
                    writer.writerow(['=== 시군구별 편입면적 분석 ==='])
                    writer.writerow(['시군구', '필지 수', '토지면적(m2)', '편입면적(m2)', '공시지가합계(원)'])
                    for sigungu, stats in sorted(dashboard_stats.get('sigungu_stats', {}).items()):
                        writer.writerow([
                            sigungu, stats['count'],
                            f"{stats['area']:,.2f}", f"{stats['inclusion_area']:,.2f}",
                            f"{stats['price_sum']:,.0f}"
                        ])
                    writer.writerow([])

                    # 지목별 통계
                    writer.writerow(['=== 지목별 편입면적 분석 ==='])
                    writer.writerow(['지목', '필지 수', '토지면적(m2)', '편입면적(m2)', '공시지가합계(원)', '비율(%)'])
                    total_count = summary.get('total_count', 0)
                    for jimok, stats in sorted(dashboard_stats.get('jimok_stats', {}).items()):
                        ratio = (stats['count'] / total_count * 100) if total_count > 0 else 0
                        writer.writerow([
                            jimok, stats['count'],
                            f"{stats['area']:,.2f}", f"{stats['inclusion_area']:,.2f}",
                            f"{stats['price_sum']:,.0f}", f"{ratio:.1f}%"
                        ])
                    writer.writerow([])

                    # 소유구분별 통계
                    writer.writerow(['=== 소유구분별 분석 ==='])
                    writer.writerow(['소유구분', '필지 수', '토지면적(m2)', '편입면적(m2)', '공시지가합계(원)'])
                    for owner_type, stats in sorted(dashboard_stats.get('owner_stats', {}).items()):
                        writer.writerow([
                            owner_type, stats['count'],
                            f"{stats['area']:,.2f}", f"{stats['inclusion_area']:,.2f}",
                            f"{stats['price_sum']:,.0f}"
                        ])
                    writer.writerow([])

                    # 공시지가 분포
                    writer.writerow(['=== 공시지가 분포 ==='])
                    writer.writerow(['가격대', '필지 수', '비율(%)'])
                    for price_dist in dashboard_stats.get('price_distribution', []):
                        writer.writerow([
                            price_dist['label'], price_dist['count'],
                            f"{price_dist['ratio']:.1f}%"
                        ])
                    writer.writerow([])

                # 연속지적도 데이터
                writer.writerow(['=== 연속지적도 ==='])
                cadastral_items = data.get('cadastral', [])
                if cadastral_items:
                    writer.writerow(['PNU', '지번', '주소', '지목', '토지/임야', '공시지가', '고시년월'])
                    for item in cadastral_items:
                        props = item.get('properties', {})
                        jibun = props.get('jibun', '')
                        pnu = props.get('pnu', '')
                        jimok = extract_jimok_from_jibun(jibun)
                        if jimok == '미분류':
                            jimok = extract_jimok_from_pnu(pnu)
                        bchk = props.get('bchk', '')
                        bchk_name = '토지' if bchk == '1' else ('임야' if bchk == '2' else bchk)
                        writer.writerow([
                            pnu, jibun, props.get('addr', ''), jimok, bchk_name,
                            props.get('jiga', ''), f"{props.get('gosi_year', '')}-{props.get('gosi_month', '')}"
                        ])
                writer.writerow([])

                # 토지임야정보 데이터
                writer.writerow(['=== 토지임야정보 ==='])
                land_forest_items = data.get('land_forest', [])
                if land_forest_items:
                    writer.writerow(['PNU', '주소', '토지구분코드', '지번'])
                    for item in land_forest_items:
                        props = item.get('properties', {})
                        writer.writerow([
                            props.get('pnu', ''), props.get('addr', ''),
                            props.get('bchk', ''), props.get('jibun', '')
                        ])
                writer.writerow([])

                # 토지특성정보 데이터
                writer.writerow(['=== 토지특성정보 ==='])
                land_char_items = data.get('land_character', [])
                if land_char_items:
                    writer.writerow(['PNU', '지번', '주소', '지목', '구분'])
                    for item in land_char_items:
                        props = item.get('properties', {})
                        jibun = props.get('jibun', '')
                        pnu = props.get('pnu', '')
                        jimok = extract_jimok_from_jibun(jibun)
                        if jimok == '미분류':
                            jimok = extract_jimok_from_pnu(pnu)
                        writer.writerow([
                            pnu, jibun, props.get('addr', ''), jimok, props.get('bchk', '')
                        ])
                writer.writerow([])

                # 개별공시지가 데이터
                writer.writerow(['=== 개별공시지가 ==='])
                land_price_items = data.get('land_price', [])
                if land_price_items:
                    writer.writerow(['PNU', '공시지가(원/m2)', '고시년도', '고시월'])
                    for item in land_price_items:
                        props = item.get('properties', {})
                        jiga = props.get('jiga', '')
                        if jiga:
                            try:
                                jiga = f"{int(float(jiga)):,}"
                            except:
                                pass
                        writer.writerow([
                            props.get('pnu', ''), jiga,
                            props.get('gosi_year', ''), props.get('gosi_month', '')
                        ])
                writer.writerow([])

                # 토지이용계획 데이터
                writer.writerow(['=== 토지이용계획 ==='])
                land_use_items = data.get('land_use_plan', [])
                if land_use_items:
                    writer.writerow(['PNU', '용도지역', '도시계획'])
                    for item in land_use_items:
                        props = item.get('properties', {})
                        writer.writerow([
                            props.get('pnu', ''),
                            props.get('prpos_area_dstrc_nm', ''),
                            props.get('cty_plan_spfc_nm', '')
                        ])
                writer.writerow([])

                # 토지소유자정보 데이터
                writer.writerow(['=== 토지소유자정보 ==='])
                land_owner_items = data.get('land_owner', [])
                if land_owner_items:
                    writer.writerow(['PNU', '주소', '지번', '소유구분', '공시지가'])
                    for item in land_owner_items:
                        props = item.get('properties', {})
                        bchk = props.get('bchk', '')
                        owner_type = '토지' if bchk == '1' else ('임야' if bchk == '2' else '기타')
                        jiga = props.get('jiga', '')
                        if jiga:
                            try:
                                jiga = f"{int(float(jiga)):,}"
                            except:
                                pass
                        writer.writerow([
                            props.get('pnu', ''), props.get('addr', ''),
                            props.get('jibun', ''), owner_type, jiga
                        ])
                writer.writerow([])

                # 건축물대장 데이터
                writer.writerow(['=== 건축물대장 ==='])
                building_items = data.get('building_register', [])
                if building_items:
                    writer.writerow(['PNU', '건물명', '동명', '주용도', '연면적', '건축면적',
                                     '건폐율', '용적률', '사용승인일', '지상층수', '지하층수'])
                    for item in building_items:
                        props = item.get('properties', {})
                        writer.writerow([
                            props.get('pnu', ''), props.get('bld_nm', ''),
                            props.get('dong_nm', ''), props.get('main_purps_cd_nm', ''),
                            props.get('tot_area', ''), props.get('arch_area', ''),
                            props.get('bc_rat', ''), props.get('vl_rat', ''),
                            props.get('use_apr_day', ''), props.get('grnd_flr_cnt', ''),
                            props.get('ugrnd_flr_cnt', '')
                        ])
                writer.writerow([])

                # 분석 탭 요약 (지형/가로경관/기반비용/AI/구역계/입지/법률/사업비)
                ExportManager._write_sections_csv(writer, analysis_summaries)

                # 디버그 로그
                if debug_log:
                    writer.writerow(['=== 디버그 로그 ==='])
                    for line in debug_log.split('\n'):
                        writer.writerow([line])

            return True
        except Exception as e:
            QgsMessageLog.logMessage(f"Excel Export Error: {e}", "VWorld", Qgis.Critical)
            return False

    @staticmethod
    def export_to_word(data, filepath, dashboard_stats=None, debug_log=None,
                       analysis_summaries=None):
        """워드 형식으로 내보내기 (전체 탭 분석 요약 포함)"""
        if filepath.lower().endswith('.docx'):
            return ExportManager._export_to_docx(
                data, filepath, dashboard_stats, debug_log, analysis_summaries)
        else:
            return ExportManager._export_to_html(
                data, filepath, dashboard_stats, debug_log, analysis_summaries)

    @staticmethod
    def _export_to_docx(data, filepath, dashboard_stats=None, debug_log=None,
                        analysis_summaries=None):
        """MS Word(.docx) 형식으로 내보내기"""
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            doc = Document()

            title = doc.add_heading('토지정보 조회 보고서', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            doc.add_paragraph(f"작성일시: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            doc.add_paragraph()

            if dashboard_stats:
                summary = dashboard_stats.get('summary', {})

                doc.add_heading('대시보드 요약', level=1)
                summary_table = doc.add_table(rows=5, cols=2)
                summary_table.style = 'Table Grid'
                summary_data = [
                    ('총 필지 수', f"{summary.get('total_count', 0):,}"),
                    ('총 토지면적(m2)', f"{summary.get('total_area', 0):,.2f}"),
                    ('총 편입면적(m2)', f"{summary.get('total_inclusion_area', 0):,.2f}"),
                    ('평균 공시지가(원/m2)', f"{summary.get('avg_price', 0):,.0f}"),
                    ('총 공시지가 합계(원)', f"{summary.get('total_price_sum', 0):,.0f}"),
                ]
                for i, (label, value) in enumerate(summary_data):
                    summary_table.rows[i].cells[0].text = label
                    summary_table.rows[i].cells[1].text = str(value)
                doc.add_paragraph()

                # 시도별 분석
                doc.add_heading('시도별 편입면적 분석', level=1)
                sido_stats = dashboard_stats.get('sido_stats', {})
                if sido_stats:
                    table = doc.add_table(rows=len(sido_stats) + 1, cols=5)
                    table.style = 'Table Grid'
                    headers = ['시도', '필지 수', '토지면적(m2)', '편입면적(m2)', '공시지가합계(원)']
                    for i, header in enumerate(headers):
                        table.rows[0].cells[i].text = header
                    for i, (sido, stats) in enumerate(sorted(sido_stats.items()), 1):
                        table.rows[i].cells[0].text = sido
                        table.rows[i].cells[1].text = f"{stats['count']:,}"
                        table.rows[i].cells[2].text = f"{stats['area']:,.2f}"
                        table.rows[i].cells[3].text = f"{stats['inclusion_area']:,.2f}"
                        table.rows[i].cells[4].text = f"{stats['price_sum']:,.0f}"
                doc.add_paragraph()

                # 시군구별 분석
                doc.add_heading('시군구별 편입면적 분석', level=1)
                sigungu_stats = dashboard_stats.get('sigungu_stats', {})
                if sigungu_stats:
                    table = doc.add_table(rows=len(sigungu_stats) + 1, cols=5)
                    table.style = 'Table Grid'
                    headers = ['시군구', '필지 수', '토지면적(m2)', '편입면적(m2)', '공시지가합계(원)']
                    for i, header in enumerate(headers):
                        table.rows[0].cells[i].text = header
                    for i, (sigungu, stats) in enumerate(sorted(sigungu_stats.items()), 1):
                        table.rows[i].cells[0].text = sigungu
                        table.rows[i].cells[1].text = f"{stats['count']:,}"
                        table.rows[i].cells[2].text = f"{stats['area']:,.2f}"
                        table.rows[i].cells[3].text = f"{stats['inclusion_area']:,.2f}"
                        table.rows[i].cells[4].text = f"{stats['price_sum']:,.0f}"
                doc.add_paragraph()

                # 지목별 분석
                doc.add_heading('지목별 편입면적 분석', level=1)
                jimok_stats = dashboard_stats.get('jimok_stats', {})
                total_count = summary.get('total_count', 0)
                if jimok_stats:
                    table = doc.add_table(rows=len(jimok_stats) + 1, cols=6)
                    table.style = 'Table Grid'
                    headers = ['지목', '필지 수', '토지면적(m2)', '편입면적(m2)', '공시지가합계(원)', '비율(%)']
                    for i, header in enumerate(headers):
                        table.rows[0].cells[i].text = header
                    for i, (jimok, stats) in enumerate(sorted(jimok_stats.items()), 1):
                        ratio = (stats['count'] / total_count * 100) if total_count > 0 else 0
                        table.rows[i].cells[0].text = jimok
                        table.rows[i].cells[1].text = f"{stats['count']:,}"
                        table.rows[i].cells[2].text = f"{stats['area']:,.2f}"
                        table.rows[i].cells[3].text = f"{stats['inclusion_area']:,.2f}"
                        table.rows[i].cells[4].text = f"{stats['price_sum']:,.0f}"
                        table.rows[i].cells[5].text = f"{ratio:.1f}%"
                doc.add_paragraph()

                # 소유구분별 분석
                doc.add_heading('소유구분별 분석', level=1)
                owner_stats = dashboard_stats.get('owner_stats', {})
                if owner_stats:
                    table = doc.add_table(rows=len(owner_stats) + 1, cols=5)
                    table.style = 'Table Grid'
                    headers = ['소유구분', '필지 수', '토지면적(m2)', '편입면적(m2)', '공시지가합계(원)']
                    for i, header in enumerate(headers):
                        table.rows[0].cells[i].text = header
                    for i, (owner_type, stats) in enumerate(sorted(owner_stats.items()), 1):
                        table.rows[i].cells[0].text = owner_type
                        table.rows[i].cells[1].text = f"{stats['count']:,}"
                        table.rows[i].cells[2].text = f"{stats['area']:,.2f}"
                        table.rows[i].cells[3].text = f"{stats['inclusion_area']:,.2f}"
                        table.rows[i].cells[4].text = f"{stats['price_sum']:,.0f}"
                doc.add_paragraph()

                # 공시지가 분포
                doc.add_heading('공시지가 분포', level=1)
                price_dist = dashboard_stats.get('price_distribution', [])
                if price_dist:
                    table = doc.add_table(rows=len(price_dist) + 1, cols=3)
                    table.style = 'Table Grid'
                    headers = ['가격대', '필지 수', '비율(%)']
                    for i, header in enumerate(headers):
                        table.rows[0].cells[i].text = header
                    for i, dist in enumerate(price_dist, 1):
                        table.rows[i].cells[0].text = dist['label']
                        table.rows[i].cells[1].text = str(dist['count'])
                        table.rows[i].cells[2].text = f"{dist['ratio']:.1f}%"
                doc.add_paragraph()

            # 연속지적도 데이터
            cadastral_items = data.get('cadastral', [])
            if cadastral_items:
                doc.add_heading(f'연속지적도 ({len(cadastral_items)}건)', level=1)
                table = doc.add_table(rows=min(len(cadastral_items), 100) + 1, cols=6)
                table.style = 'Table Grid'
                headers = ['PNU', '지번', '주소', '지목', '토지/임야', '공시지가']
                for i, header in enumerate(headers):
                    table.rows[0].cells[i].text = header
                for i, item in enumerate(cadastral_items[:100], 1):
                    props = item.get('properties', {})
                    jibun = props.get('jibun', '')
                    pnu = props.get('pnu', '')
                    jimok = extract_jimok_from_jibun(jibun)
                    if jimok == '미분류':
                        jimok = extract_jimok_from_pnu(pnu)
                    bchk = props.get('bchk', '')
                    bchk_name = '토지' if bchk == '1' else ('임야' if bchk == '2' else bchk)
                    jiga = props.get('jiga', '')
                    if jiga:
                        try:
                            jiga = f"{int(float(jiga)):,}"
                        except:
                            pass
                    table.rows[i].cells[0].text = str(pnu)
                    table.rows[i].cells[1].text = str(jibun)
                    table.rows[i].cells[2].text = str(props.get('addr', ''))
                    table.rows[i].cells[3].text = jimok
                    table.rows[i].cells[4].text = bchk_name
                    table.rows[i].cells[5].text = str(jiga)
                doc.add_paragraph()

            # 토지이용계획 데이터
            land_use_items = data.get('land_use_plan', [])
            if land_use_items:
                doc.add_heading(f'토지이용계획 ({len(land_use_items)}건)', level=1)
                table = doc.add_table(rows=min(len(land_use_items), 100) + 1, cols=3)
                table.style = 'Table Grid'
                headers = ['PNU', '용도지역', '도시계획']
                for i, header in enumerate(headers):
                    table.rows[0].cells[i].text = header
                for i, item in enumerate(land_use_items[:100], 1):
                    props = item.get('properties', {})
                    table.rows[i].cells[0].text = str(props.get('pnu', ''))
                    table.rows[i].cells[1].text = str(props.get('prpos_area_dstrc_nm', ''))
                    table.rows[i].cells[2].text = str(props.get('cty_plan_spfc_nm', ''))
                doc.add_paragraph()

            # 건축물대장 데이터
            building_items = data.get('building_register', [])
            if building_items:
                doc.add_heading(f'건축물대장 ({len(building_items)}건)', level=1)
                table = doc.add_table(rows=min(len(building_items), 100) + 1, cols=6)
                table.style = 'Table Grid'
                headers = ['PNU', '건물명', '주용도', '연면적', '사용승인일', '층수']
                for i, header in enumerate(headers):
                    table.rows[0].cells[i].text = header
                for i, item in enumerate(building_items[:100], 1):
                    props = item.get('properties', {})
                    table.rows[i].cells[0].text = str(props.get('pnu', ''))
                    table.rows[i].cells[1].text = str(props.get('bld_nm', ''))
                    table.rows[i].cells[2].text = str(props.get('main_purps_cd_nm', ''))
                    table.rows[i].cells[3].text = str(props.get('tot_area', ''))
                    table.rows[i].cells[4].text = str(props.get('use_apr_day', ''))
                    grnd = props.get('grnd_flr_cnt', '')
                    ugrnd = props.get('ugrnd_flr_cnt', '')
                    table.rows[i].cells[5].text = f"지상{grnd}/지하{ugrnd}"
                doc.add_paragraph()

            # 분석 탭 요약 (지형/가로경관/기반비용/AI/구역계/입지/법률/사업비)
            ExportManager._write_sections_docx(doc, analysis_summaries)

            # 디버그 로그
            if debug_log:
                doc.add_heading('디버그 로그', level=1)
                doc.add_paragraph(debug_log[:5000])

            doc.save(filepath)
            return True

        except ImportError:
            QgsMessageLog.logMessage("python-docx 라이브러리가 설치되어 있지 않습니다. HTML 형식으로 저장합니다.", "VWorld", Qgis.Warning)
            html_filepath = filepath.replace('.docx', '.html')
            return ExportManager._export_to_html(
                data, html_filepath, dashboard_stats, debug_log,
                analysis_summaries)
        except Exception as e:
            QgsMessageLog.logMessage(f"DOCX Export Error: {e}", "VWorld", Qgis.Critical)
            return False

    @staticmethod
    def _export_to_html(data, filepath, dashboard_stats=None, debug_log=None,
                        analysis_summaries=None):
        """HTML 형식으로 내보내기"""
        try:
            html_content = """
<!DOCTYPE html>
<html>
<head>
    <meta charset="utf-8">
    <title>토지정보 조회 보고서</title>
    <style>
        body { font-family: 'Malgun Gothic', sans-serif; margin: 20px; }
        h1 { color: #2c3e50; border-bottom: 2px solid #3498db; padding-bottom: 10px; }
        h2 { color: #34495e; margin-top: 30px; border-left: 4px solid #3498db; padding-left: 10px; }
        table { border-collapse: collapse; width: 100%; margin: 10px 0; }
        th, td { border: 1px solid #bdc3c7; padding: 8px; text-align: left; }
        th { background-color: #3498db; color: white; }
        tr:nth-child(even) { background-color: #ecf0f1; }
        .summary { background-color: #e8f6f3; padding: 15px; border-radius: 5px; margin: 20px 0; }
        .stats-grid { display: grid; grid-template-columns: repeat(3, 1fr); gap: 10px; }
        .stat-item { background: #f8f9fa; padding: 10px; border-radius: 5px; text-align: center; }
        .stat-value { font-size: 24px; font-weight: bold; color: #2c3e50; }
        .stat-label { font-size: 12px; color: #7f8c8d; }
        .debug-log { background-color: #f5f5f5; padding: 10px; border: 1px solid #ddd; font-family: monospace; font-size: 11px; white-space: pre-wrap; max-height: 500px; overflow-y: auto; }
    </style>
</head>
<body>
    <h1>토지정보 조회 보고서</h1>
    <p>작성일시: """ + datetime.now().strftime("%Y-%m-%d %H:%M:%S") + """</p>
"""

            if dashboard_stats:
                summary = dashboard_stats.get('summary', {})
                html_content += """
    <div class="summary">
        <h2>조회 요약</h2>
        <div class="stats-grid">
            <div class="stat-item">
                <div class="stat-value">""" + f"{summary.get('total_count', 0):,}" + """</div>
                <div class="stat-label">총 필지 수</div>
            </div>
            <div class="stat-item">
                <div class="stat-value">""" + f"{summary.get('total_area', 0):,.2f}" + """</div>
                <div class="stat-label">총 토지면적(m2)</div>
            </div>
            <div class="stat-item">
                <div class="stat-value">""" + f"{summary.get('total_inclusion_area', 0):,.2f}" + """</div>
                <div class="stat-label">총 편입면적(m2)</div>
            </div>
            <div class="stat-item">
                <div class="stat-value">""" + f"{summary.get('avg_price', 0):,.0f}" + """</div>
                <div class="stat-label">평균 공시지가(원/m2)</div>
            </div>
            <div class="stat-item">
                <div class="stat-value">""" + f"{summary.get('total_price_sum', 0):,.0f}" + """</div>
                <div class="stat-label">총 공시지가 합계(원)</div>
            </div>
            <div class="stat-item">
                <div class="stat-value">""" + f"{summary.get('owner_count', 0):,}" + """</div>
                <div class="stat-label">소유자 수</div>
            </div>
        </div>
    </div>
"""
                # 시도별 분석
                html_content += """
    <h2>시도별 편입면적 분석</h2>
    <table>
        <tr><th>시도</th><th>필지 수</th><th>토지면적(m2)</th><th>편입면적(m2)</th><th>공시지가합계(원)</th></tr>
"""
                for sido, stats in sorted(dashboard_stats.get('sido_stats', {}).items()):
                    html_content += f"        <tr><td>{sido}</td><td>{stats['count']:,}</td><td>{stats['area']:,.2f}</td><td>{stats['inclusion_area']:,.2f}</td><td>{stats['price_sum']:,.0f}</td></tr>\n"
                html_content += "    </table>\n"

                # 시군구별 분석
                html_content += """
    <h2>시군구별 편입면적 분석</h2>
    <table>
        <tr><th>시군구</th><th>필지 수</th><th>토지면적(m2)</th><th>편입면적(m2)</th><th>공시지가합계(원)</th></tr>
"""
                for sigungu, stats in sorted(dashboard_stats.get('sigungu_stats', {}).items()):
                    html_content += f"        <tr><td>{sigungu}</td><td>{stats['count']:,}</td><td>{stats['area']:,.2f}</td><td>{stats['inclusion_area']:,.2f}</td><td>{stats['price_sum']:,.0f}</td></tr>\n"
                html_content += "    </table>\n"

                # 지목별 분석
                total_count = summary.get('total_count', 0)
                html_content += """
    <h2>지목별 편입면적 분석</h2>
    <table>
        <tr><th>지목</th><th>필지 수</th><th>토지면적(m2)</th><th>편입면적(m2)</th><th>공시지가합계(원)</th><th>비율(%)</th></tr>
"""
                for jimok, stats in sorted(dashboard_stats.get('jimok_stats', {}).items()):
                    ratio = (stats['count'] / total_count * 100) if total_count > 0 else 0
                    html_content += f"        <tr><td>{jimok}</td><td>{stats['count']:,}</td><td>{stats['area']:,.2f}</td><td>{stats['inclusion_area']:,.2f}</td><td>{stats['price_sum']:,.0f}</td><td>{ratio:.1f}%</td></tr>\n"
                html_content += "    </table>\n"

                # 소유구분별 분석
                html_content += """
    <h2>소유구분별 분석</h2>
    <table>
        <tr><th>소유구분</th><th>필지 수</th><th>토지면적(m2)</th><th>편입면적(m2)</th><th>공시지가합계(원)</th></tr>
"""
                for owner_type, stats in sorted(dashboard_stats.get('owner_stats', {}).items()):
                    html_content += f"        <tr><td>{owner_type}</td><td>{stats['count']:,}</td><td>{stats['area']:,.2f}</td><td>{stats['inclusion_area']:,.2f}</td><td>{stats['price_sum']:,.0f}</td></tr>\n"
                html_content += "    </table>\n"

                # 공시지가 분포
                html_content += """
    <h2>공시지가 분포</h2>
    <table>
        <tr><th>가격대</th><th>필지 수</th><th>비율(%)</th></tr>
"""
                for dist in dashboard_stats.get('price_distribution', []):
                    html_content += f"        <tr><td>{dist['label']}</td><td>{dist['count']:,}</td><td>{dist['ratio']:.1f}%</td></tr>\n"
                html_content += "    </table>\n"

            # 연속지적도 데이터
            cadastral_items = data.get('cadastral', [])
            if cadastral_items:
                html_content += f"\n    <h2>연속지적도 ({len(cadastral_items)}건)</h2>\n"
                html_content += "    <table>\n        <tr><th>PNU</th><th>지번</th><th>주소</th><th>지목</th><th>토지/임야</th><th>공시지가</th><th>고시년월</th></tr>\n"
                for item in cadastral_items[:100]:
                    props = item.get('properties', {})
                    jibun = props.get('jibun', '')
                    pnu = props.get('pnu', '')
                    jimok = extract_jimok_from_jibun(jibun)
                    if jimok == '미분류':
                        jimok = extract_jimok_from_pnu(pnu)
                    bchk = props.get('bchk', '')
                    bchk_name = '토지' if bchk == '1' else ('임야' if bchk == '2' else bchk)
                    jiga = props.get('jiga', '')
                    if jiga:
                        try:
                            jiga = f"{int(float(jiga)):,}"
                        except:
                            pass
                    html_content += f"        <tr><td>{pnu}</td><td>{jibun}</td><td>{props.get('addr', '')}</td><td>{jimok}</td><td>{bchk_name}</td><td>{jiga}</td><td>{props.get('gosi_year', '')}-{props.get('gosi_month', '')}</td></tr>\n"
                html_content += "    </table>\n"

            # 분석 탭 요약 (지형/가로경관/기반비용/AI/구역계/입지/법률/사업비)
            html_content += ExportManager._render_sections_html(
                analysis_summaries)

            # 디버그 로그
            if debug_log:
                html_content += '\n    <h2>디버그 로그</h2>\n    <div class="debug-log">\n'
                html_content += debug_log[:10000]
                html_content += "\n    </div>\n"

            html_content += "\n</body>\n</html>\n"

            with open(filepath, 'w', encoding='utf-8') as f:
                f.write(html_content)

            return True
        except Exception as e:
            QgsMessageLog.logMessage(f"HTML Export Error: {e}", "VWorld", Qgis.Critical)
            return False
